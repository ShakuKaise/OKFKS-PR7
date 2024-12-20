import zipfile
from datetime import date, timedelta
import pandas as pd

from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import LogoutView
from django.contrib.auth.views import LoginView as DjangoLoginView  # Переименовали
from django.core.files.storage import FileSystemStorage
from django.db.models import Count
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse, HttpResponseForbidden
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.views import View
from django.views.decorators.http import require_POST
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView

from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.pagination import PageNumberPagination
from rest_framework.views import APIView
from rest_framework.authtoken.models import Token
from rest_framework.response import Response
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from drf_spectacular.utils import extend_schema
import io
import os
import csv
import subprocess
from django.db import connection
from django.conf import settings
from .models import *
from .forms import *
from .serializers import (UserSerializer, BookSerializer, RequestSerializer, PublisherSerializer, AuthorSerializer,
                          GenreSerializer, LanguageSerializer)

# Authentication Views
class RegisterView(View):
    template_name = 'reg/registration.html'

    def get(self, request):
        form = UserForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request):
        form = UserForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password'])  # Hash the password
            user.save()
            return redirect('profile')
        return render(request, self.template_name, {'form': form})


class LoginView(APIView):
    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(username=username, password=password)
        if user is not None:
            token, created = Token.objects.get_or_create(user=user)
            return Response({'token': token.key}, status=status.HTTP_200_OK)
        return Response({'error': 'Invalid Credentials'}, status=status.HTTP_400_BAD_REQUEST)


class CustomLoginView(DjangoLoginView):
    template_name = 'reg/login.html'
    authentication_form = AuthenticationForm
    next_page = 'home'


def logout_view(request):
    logout(request)
    return redirect('home')

def home_view(request):
    books = Book.objects.filter(is_deleted=False)
    user = request.user
    return render(request, 'Books/book_list.html', {'books': books, 'user': user})

@login_required
def profile(request):
    user = request.user

    # Извлекаем запросы на аренду с их книгами
    user_requests = Request.objects.filter(user=user, status=Request.RequestStatus.RENTED)

    # Получаем список арендованных книг
    rented_books = [req.book for req in user_requests]

    return render(request, 'profile.html', {
        'user': user,
        'rented_books': rented_books,
        'requests': user_requests,
    })  # 'user': user is already included here.

@login_required
def rented_books_statistics(request):

    # 1. Получаем данные о возвращенных книгах
    returned_requests = (
        Request.objects.all()
        .values('book__name')
        .annotate(count=Count('id'))
    )
    # Преобразуем данные в удобный формат
    book_names = [req['book__name'] for req in returned_requests]
    return_counts = [req['count'] for req in returned_requests]

    plt.figure(figsize=(10, 6))
    plt.pie(return_counts, labels=book_names, autopct='%1.1f%%', startangle=90, colors=plt.cm.tab20.colors)
    plt.title('Статистика аренды книг')
    plt.axis('equal')  # Обеспечиваем круговую диаграмму

    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    plt.close()

    document = Document()
    document.add_heading('Статистика когда-либо арендованных книг', level=1)
    document.add_paragraph('Диаграмма ниже отображает статистику по книгам, которые когда-либо были арендованны:')

    # Вставляем диаграмму в Word
    image_path = 'returned_books_chart.png'
    with open(image_path, 'wb') as image_file:
        image_file.write(buffer.getvalue())
    document.add_picture(image_path, width=Inches(5))
    os.remove(image_path)

    # Добавляем текстовый список после диаграммы
    document.add_paragraph('Список статистики по книгам:')
    for name, count in zip(book_names, return_counts):
        document.add_paragraph(f"- {name}: {count} раз(а)")

    response_buffer = BytesIO()
    document.save(response_buffer)
    response_buffer.seek(0)

    response = HttpResponse(response_buffer,
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="returned_books_statistics.docx"'
    return response

def return_book(request, book_id):
    # Найти запись о книге по ID и объекту пользователя
    rented_book = get_object_or_404(Request, id=book_id, user=request.user, status=Request.RequestStatus.RENTED)

    # Выполнить проверку, если это обязательно (например, только владелец может завершить аренду)
    if rented_book.user != request.user:
        return HttpResponseForbidden("Вы не можете вернуть эту книгу.")

    # Обновить статус книги на "вернулась"
    rented_book.status = Request.RequestStatus.RETURNED
    rented_book.save()

    # Добавьте перенаправление на страницу «Профиль» или другую соответствующую страницу
    return redirect('profile')

def rent_book(request, book_id):
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Вы должны быть авторизованы, чтобы арендовать книгу.'}, status=401)

    book = get_object_or_404(Book, id=book_id, is_deleted=False)

    # Проверка, арендовал ли пользователь эту книгу
    already_rented = Request.objects.filter(
        book=book,
        user=request.user,
        status=Request.RequestStatus.RENTED
    ).exists()

    if already_rented:
        return JsonResponse({'error': 'Вы уже арендовали эту книгу.'}, status=400)

    # Создание записи об аренде
    Request.objects.create(
        user=request.user,
        book=book,
        borrow_date=date.today(),
        return_date=date.today() + timedelta(days=14),
        status=Request.RequestStatus.RENTED
    )

    return

def backup_database(request):
    """
    Создать резервную копию PostgreSQL в SQL формате и скачать файл
    """
    db_settings = settings.DATABASES['default']
    db_name = db_settings['NAME']
    db_user = db_settings['USER']
    db_password = db_settings['PASSWORD']  # Получить пароль пользователя PostgreSQL
    db_host = db_settings.get('HOST', 'localhost')
    db_port = db_settings.get('PORT', 5432)

    try:
        # Создаем буфер в памяти
        buffer = io.BytesIO()

        # Устанавливаем переменную окружения для пароля PostgreSQL
        env = os.environ.copy()
        env['PGPASSWORD'] = db_password

        # Выполняем `pg_dump` и записываем вывод в буфер
        process = subprocess.Popen(
            [
                'C:\\Program Files\\PostgreSQL\\16\\bin\\pg_dump.exe',
                '-h', db_host,
                '-p', str(db_port),
                '-U', db_user,
                '-F', 'p',  # Формат plain text
                db_name
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=False,  # Считываем байтовые данные
            env=env  # Передаем окружение с паролем
        )
        stdout, stderr = process.communicate()

        # Проверяем ошибки
        if process.returncode != 0:
            return HttpResponse(f"Ошибка при создании резервной копии: {stderr.decode('utf-8')}", status=500)

        # Пишем результат в буфер
        buffer.write(stdout)
        buffer.seek(0)

        # Настраиваем HTTP-ответ для скачивания
        response = HttpResponse(buffer, content_type='application/sql')
        response['Content-Disposition'] = f'attachment; filename="{db_name}_backup.sql"'
        return response

    except Exception as e:
        return HttpResponse(f"Ошибка: {str(e)}", status=500)

def upload_file_view(request):
    if request.method == "POST":
        form = FileUploadForm(request.POST, request.FILES)  # Используем request.FILES для получения загруженного файла
        if form.is_valid():
            uploaded_file = request.FILES['file']

            # Сохраняем файл в папку "uploads" в корневой директории проекта
            upload_path = os.path.join(settings.BASE_DIR, 'uploads', uploaded_file.name)
            os.makedirs(os.path.dirname(upload_path), exist_ok=True)

            with open(upload_path, 'wb') as f:
                for chunk in uploaded_file.chunks():
                    f.write(chunk)

            # Можно сразу вызывать обработку файла (например, импорт данных в базу)
            return HttpResponse(f"Файл {uploaded_file.name} успешно загружен и сохранен.")
    else:
        form = FileUploadForm()

    user = request.user
    return render(request, 'import_backup.html', {'form': form, 'user': user})

def restore_database(request):
    """
    Восстановление базы данных из загруженной резервной копии SQL
    """
    if request.method == 'GET':
        # При GET запросе отображаем форму
        form = FileUploadForm()
        return render(request, 'import_backup.html', {'form': form})

    # При POST запросе обрабатываем загруженный файл
    form = FileUploadForm(request.POST, request.FILES)

    if not form.is_valid():
        return HttpResponse("Неверная форма. Убедитесь, что прикреплен файл.", status=400)

    # Достаём файл из формы
    uploaded_file = request.FILES['file']

    # Параметры базы данных из Django settings
    db_settings = settings.DATABASES['default']
    db_name = db_settings['NAME']
    db_user = db_settings['USER']
    db_password = db_settings['PASSWORD']
    db_host = db_settings.get('HOST', 'localhost')
    db_port = db_settings.get('PORT', 5432)

    try:
        # Устанавливаем переменную окружения для пароля PostgreSQL
        env = os.environ.copy()
        env['PGPASSWORD'] = db_password

        # Выполняем команду psql для восстановления базы данслоаных
        subprocess.run(
            [
                r'C:\Program Files\PostgreSQL\16\bin\psql.exe',  # Добавьте полный путь к psql
                '-h', db_host,
                '-p', str(db_port),
                '-U', db_user,
                '-d', db_name
            ],
            input=uploaded_file.read().decode('utf-8'),
            text=True,
            env=env,
            check=True
        )

        return HttpResponse("База данных успешно восстановлена из загруженной резервной копии.")
    except subprocess.CalledProcessError as e:
        return HttpResponse(f"Ошибка при восстановлении базы данных: {e.stderr}", status=500)
    except Exception as e:
        return HttpResponse(f"Ошибка: {str(e)}", status=500)

def restore_db(request):
    """Restore the database from an uploaded SQL file."""
    if request.method == 'POST' and request.FILES.get('backup_file'):
        backup_file = request.FILES['backup_file']
        backup_format = os.path.splitext(backup_file.name)[-1].lower()

        if backup_format != '.sql':
            return redirect('custom_admin:dashboard')

        try:
            # Save the uploaded file to a temporary folder
            fs = FileSystemStorage(location='temp_backups')
            os.makedirs('temp_backups', exist_ok=True)
            file_path = fs.save(backup_file.name, backup_file)
            full_path = os.path.join('temp_backups', file_path)

            # Execute the SQL commands using the PostgreSQL connection
            with connection.cursor() as cursor:
                with open(full_path, 'r') as sql_file:
                    sql_content = sql_file.read()
                    cursor.execute(sql_content)

            # Cleanup temporary backup file
            fs.delete(file_path)

            return redirect('custom_admin:dashboard')
        except Exception as e:
            return redirect('custom_admin:dashboard')

    return redirect('custom_admin:dashboard')

# CRUD Views

class AuthorCreateView(CreateView):
    model = Author
    template_name = 'Authors/author_create.html'
    form_class = AuthorForm
    success_url = reverse_lazy('authors')


class AuthorsList(ListView):
    model = Author
    template_name = 'Authors/author_list.html'
    context_object_name = 'list_authors'
    paginate_by = 12

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['authors'] = Author.objects.filter(is_deleted=False)
        context['title'] = 'Авторы'
        return context


class AuthorUpdateView(UpdateView):
    model = Author
    form_class = AuthorForm
    template_name = 'Authors/author_create.html'  # Updated to use create.html template
    success_url = reverse_lazy('authors')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = "Edit Author"
        context['button_text'] = "Update"
        return context


class AuthorDeleteView(DeleteView):
    model = Author
    success_url = reverse_lazy('authors')

    def dispatch(self, request, *args, **kwargs):
        if request.method == "POST":
            author = self.get_object()
            author.is_deleted = True
            author.save()
            return redirect(self.success_url)
        else:
            return super().dispatch(request, *args, **kwargs)

class GenreUpdateView(UpdateView):
    model = Genre
    form_class = GenreForm
    template_name = 'Genres/genre_create.html'  # Используем существующий шаблон создания
    success_url = reverse_lazy('genres')  # После успешного редактирования - назад в список

    def get_context_data(self, **kwargs):
        # Добавляем динамический заголовок в шаблон
        context = super().get_context_data(**kwargs)
        context['title'] = "Edit Genre"  # Название страницы
        context['button_text'] = "Update"  # Текст кнопки
        return context


class GenreDeleteView(DeleteView):
    model = Genre
    success_url = reverse_lazy('genres')  # После удаления жанра возвращаемся на список

    def dispatch(self, request, *args, **kwargs):
        # Проверяем, если запрос типа POST, сразу выполняем удаление
        if request.method == "POST":
            genre = self.get_object()
            genre.is_deleted = True  # Или genre.delete(), если вам нужно реальное удаление
            genre.save()
            return redirect(self.success_url)
        else:
            return super().dispatch(request, *args, **kwargs)


class BookUpdateView(UpdateView):
    model = Book
    template_name = 'Books/book_create.html'  # Updated to use create.html template
    form_class = BookForm
    success_url = reverse_lazy('book_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = "Edit Book"
        context['button_text'] = "Update"
        return context


class BookDeleteView(DeleteView):
    model = Book
    success_url = reverse_lazy('book_list')

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.object.is_deleted = True
        self.object.save()
        return HttpResponseRedirect(self.success_url)


class PublisherCreateView(CreateView):
    model = Publisher
    template_name = 'Publishers/publisher_create.html'
    form_class = PublisherForm
    success_url = reverse_lazy('publishers')


class PublisherUpdateView(UpdateView):
    model = Publisher
    template_name = 'Publishers/publisher_create.html'  # Updated to use create.html template
    form_class = PublisherForm
    success_url = reverse_lazy('publishers')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = "Edit Publisher"
        context['button_text'] = "Update"
        return context


class PublisherDeleteView(DeleteView):
    model = Publisher
    success_url = reverse_lazy('publishers')

    def dispatch(self, request, *args, **kwargs):
        if request.method == "POST":
            publisher = self.get_object()
            publisher.is_deleted = True
            publisher.save()
            return redirect(self.success_url)
        else:
            return super().dispatch(request, *args, **kwargs)


class LanguageCreateView(CreateView):
    model = Language
    template_name = 'Languages/language_create.html'
    form_class = LanguageForm
    success_url = reverse_lazy('languages')



class LanguageUpdateView(UpdateView):
    model = Language
    template_name = 'Languages/language_create.html'  # Updated to use create.html template
    form_class = LanguageForm
    success_url = reverse_lazy('languages')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = "Edit Language"
        context['button_text'] = "Update"
        return context


class LanguageDeleteView(DeleteView):
    model = Language
    success_url = reverse_lazy('languages')

    def dispatch(self, request, *args, **kwargs):
        if request.method == "POST":
            language = self.get_object()
            language.is_deleted = True
            language.save()
            return redirect(self.success_url)
        else:
            return super().dispatch(request, *args, **kwargs)


class RequestCreateView(CreateView):
    model = Request
    template_name = 'Requests/request_create.html'
    form_class = RequestForm
    success_url = reverse_lazy('requests')


class RequestListView(ListView):
    model = Request
    template_name = 'Requests/request_list.html'
    context_object_name = 'list_requests'
    paginate_by = 12


class RequestUpdateView(UpdateView):
    model = Request
    template_name = 'Requests/request_create.html'  # Updated to use create.html template
    form_class = RequestForm
    success_url = reverse_lazy('requests')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = "Edit Request"
        context['button_text'] = "Update"
        return context


class RequestDeleteView(DeleteView):
    model = Request
    success_url = reverse_lazy('requests')

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.object.is_deleted = True
        self.object.save()
        return HttpResponseRedirect(self.success_url)


class AuthorsList(ListView):
    model = Author
    template_name = 'Authors/author_list.html'
    context_object_name = 'list_authors'
    paginate_by = 12

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['authors'] = Author.objects.filter(is_deleted=False)
        context['title'] = 'Авторы'
        return context


class GenreCreateView(CreateView):
    model = Genre
    template_name = 'Genres/genre_create.html'
    form_class = GenreForm
    success_url = reverse_lazy('genres')

class GenresList(ListView):
    model = Genre
    template_name = 'Genres/genre_list.html'
    context_object_name = 'list_genres'
    paginate_by = 12

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['genres'] = Genre.objects.filter(is_deleted=False)
        context['title'] = 'Жанры'
        return context

class LanguagesList(ListView):
    model = Language
    template_name = 'Languages/language_list.html'
    context_object_name = 'list_languages'
    paginate_by = 12

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['languages'] = Language.objects.filter(is_deleted=False)
        context['title'] = 'Языки'
        return context

class PublishersList(ListView):
    model = Publisher
    template_name = 'Publishers/publisher_list.html'
    context_object_name = 'list_publishers'
    paginate_by = 12

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['publishers'] = Publisher.objects.filter(is_deleted=False)
        context['title'] = 'Издатели'
        return context

class RequestsList(ListView):
    model = Request
    template_name = 'Requests/request_list.html'
    context_object_name = 'list_requests'
    paginate_by = 12

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['requests'] = Request.objects.all()
        context['title'] = 'Запросы'
        return context


class BookCreateView(CreateView):
    model = Book
    template_name = 'Books/book_create.html'
    form_class = BookForm
    success_url = reverse_lazy('book_list')


class BookList(ListView):
    model = Book
    template_name = 'Books/book_list.html'
    context_object_name = 'list_books'
    paginate_by = 2

    def get_queryset(self):
        queryset = Book.objects.filter(is_deleted=False)
        form = BookFilterForm(self.request.GET)

        if form.is_valid():
            tag = form.cleaned_data.get('tag')
            category = form.cleaned_data.get('category')

            if category:
                queryset = queryset.filter(category=category)
            if tag:
                queryset = queryset.filter(tags=tag)

        return queryset

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data()
        context['form'] = BookFilterForm(self.request.GET)
        return context


# Utility Functions
def index(request):
    return render(request, 'index.html')

def book_list(request):
    books = Book.objects.filter(is_deleted=False)
    genres = Genre.objects.all()
    authors = Author.objects.all()
    user = request.user
    selected_genre = request.GET.get('genre')
    selected_author = request.GET.get('author')
    title_query = request.GET.get('title', '').strip()

    if selected_genre:
        books = books.filter(genres__id=selected_genre)

    if selected_author:
        books = books.filter(authors__id=selected_author)

    if title_query:
        books = books.filter(name__icontains=title_query)

    return render(request, 'Books/book_list.html', {
        'books': books,
        'genres': genres,
        'authors': authors,
        'selected_genre': selected_genre,
        'selected_author': selected_author,
        'user': user,
    })


class CustomPagination(PageNumberPagination):
    page_size = 1  # Количество объектов на странице
    page_size_query_param = 'page_size'
    max_page_size = 1000  # Максимальное количество объектов на странице

def user_rented_books_view(request):
    # Проверяем, что пользователь авторизован
    if not request.user.is_authenticated:
        return redirect('login')

    # Фильтруем запросы по текущему пользователю и статусу 'RENTED'
    rented_books = Request.objects.filter(
        user=request.user,
        status=Request.RequestStatus.RENTED
    ).select_related('book')

    # Передаем арендованные книги в шаблон
    context = {
        'rented_books': rented_books
    }
    return render(request, 'path_to_template.html', context)

class UserViewSet(viewsets.ModelViewSet):
    serializer_class = UserSerializer
    queryset = User.objects.all()
    pagination_class = CustomPagination


class CustomModelViewSet(viewsets.ModelViewSet):
    @action(detail=False, methods=['patch'], url_path='delete-multiple')
    def delete_multiple(self, request):
        """
        Метод для множественного логического удаления объектов.

        Ожидается JSON-объект вида {"ids": [1, 2, 3, ...]}
        где ids - список идентификаторов объектов для удаления.
        """
        ids = request.data.get('ids', [])
        queryset = self.get_queryset().filter(id__in=ids)

        if not queryset.exists():
            return Response("No objects found with the provided IDs.", status=status.HTTP_404_NOT_FOUND)

        queryset.update(is_deleted=True)
        return Response("Objects successfully marked as deleted.", status=status.HTTP_200_OK)

    @action(detail=False, methods=['patch'], url_path='restore-multiple')
    def restore_multiple(self, request):
        """
        Метод для множественного восстановления объектов.

        Ожидается JSON-объект вида {"ids": [1, 2, 3, ...]}
        где ids - список идентификаторов объектов для восстановления.
        """
        ids = request.data.get('ids', [])
        queryset = self.get_queryset().filter(id__in=ids)

        if not queryset.exists():
            return Response("No objects found with the provided IDs.", status=status.HTTP_404_NOT_FOUND)

        queryset.update(is_deleted=False)
        return Response("Objects successfully restored.", status=status.HTTP_200_OK)


@extend_schema(tags=['Книги'])
class BookViewSet(CustomModelViewSet):
    queryset = Book.objects.all()
    serializer_class = BookSerializer

    def get_queryset(self):
        return self.queryset.filter(is_deleted=False)

@extend_schema(tags=['Жанры'])
class GenreViewSet(CustomModelViewSet):
    queryset = Genre.objects.all()
    serializer_class = GenreSerializer

    def get_queryset(self):
        return self.queryset.filter(is_deleted=False)

@extend_schema(tags=['Авторы'])
class AuthorViewSet(CustomModelViewSet):
    queryset = Author.objects.all()
    serializer_class = AuthorSerializer

    def get_queryset(self):
        return self.queryset.filter(is_deleted=False)

@extend_schema(tags=['Запросы'])
class RequestViewSet(viewsets.ModelViewSet):
    serializer_class = RequestSerializer
    queryset = Request.objects.all()
    pagination_class = CustomPagination


@extend_schema(tags=['Языки'])
class LanguageViewSet(viewsets.ModelViewSet):
    serializer_class = LanguageSerializer
    queryset = Language.objects.all()
    pagination_class = CustomPagination


@extend_schema(tags=['Издатели'])
class PublisherViewSet(viewsets.ModelViewSet):
    serializer_class = PublisherSerializer
    queryset = Genre.objects.all()
    pagination_class = CustomPagination


def api(request):
    return render(request, 'api.html')


def export_books_to_excel(request):
    books = Book.objects.all()
    data = {
        'Title': [book.name for book in books],
        'Publication Year': [book.publication_year for book in books],
        'Authors': [", ".join(f"{author.first_name} {author.last_name}" for author in book.authors.all()) for book in books],
        'Genres': [", ".join(genre.name for genre in book.genres.all()) for book in books],
    }
    df = pd.DataFrame(data)
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="books.xlsx"'
    df.to_excel(response, index=False)
    return response
